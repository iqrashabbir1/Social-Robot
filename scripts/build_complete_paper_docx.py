from __future__ import annotations

import csv
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
FINAL_PACKAGE = PROJECT_ROOT / "docs" / "paper1" / "final_package"
COMPANION_MD = FINAL_PACKAGE / "manuscript_sections_4_6_with_plots.md"
OUTPUT_DOCX = FINAL_PACKAGE / "paemdt_complete_updated_manuscript_zone_aligned.docx"
OUTPUT_MD = FINAL_PACKAGE / "paemdt_complete_updated_manuscript_zone_aligned.md"


TITLE = "PAEMDT: Privacy-Aware Explainable Multimodal Digital-Twin Cognitive Caregiving Robot"
AUTHORS = "Iqra Shabbir"

ABSTRACT = (
    "This paper presents PAEMDT, a privacy-aware explainable multimodal digital-twin "
    "framework for cognitive caregiving robots, with emphasis on cross-domain emotion "
    "recognition, synchronization, robustness, and deployment-aware validation. The "
    "enhanced evaluation combines source-domain benchmarking, cross-domain adaptation, "
    "repeated cross-validation, calibration analysis, missing-modality stress testing, "
    "differential privacy, digital-twin predictive validation, and edge deployment "
    "profiling. On RAVDESS, validation performance remains high (97.81% for the "
    "source-only baseline and 95.12% for the privacy-enhanced adapted model), while "
    "CREMA-D external accuracy improves from 28.30% in the source-only setting to "
    "64.28% with gradient reversal, multi-kernel MMD, and progressive pseudo-labeling, "
    "and remains 62.15% under DP-SGD at epsilon = 2.3. The measured domain gap is "
    "therefore reduced from 69.51 to 32.63 percentage points. Calibration improves to "
    "ECE = 0.041, digital-twin synchronization is measured at 124.0 +/- 67.0 ms, and "
    "Raspberry Pi 4 edge inference satisfies the real-time constraint at 47.3 ms "
    "latency (approximately 21 FPS). These results constitute technical and experimental "
    "validation rather than clinical deployment evidence, but they establish a reproducible, "
    "privacy-aware, and deployment-conscious foundation for future cognitive "
    "caregiving robot studies."
)

INTRODUCTION = [
    (
        "Cognitive caregiving robots require more than perception models. They also "
        "require a measurable runtime backbone that can synchronize sensing, mirror "
        "system state in a digital twin, and support controlled replay, audit, and "
        "disturbance analysis. This repository originally contained baseline social-robot "
        "perception code, but not a clean first-paper experimental path. PAEMDT addresses "
        "that gap by reorganizing the codebase around three foundational layers: "
        "digital-twin validation, multimodal synchronization, and emotion-recognition "
        "benchmarking."
    ),
    (
        "The resulting contribution is a technical and publication-oriented framework "
        "rather than a deployment claim. The repository preserves an implemented visual "
        "baseline, adds reproducible evaluation infrastructure, and separates synthetic, "
        "pilot, replay-grounded, and external-dataset evidence so that the maturity of "
        "each result remains explicit."
    ),
    (
        "A key enhancement of the revised manuscript is the introduction of domain "
        "adaptation to address the original cross-corpus generalization gap between "
        "RAVDESS and CREMA-D. This addition allows the paper to move from a strong "
        "source-domain benchmark to a more credible translational research platform."
    ),
]

FRAMEWORK_ARCHITECTURE = [
    (
        "PAEMDT is organized as a modular multimodal architecture in which sensing, "
        "synchronization, benchmarking, privacy controls, and digital-twin services are "
        "explicitly separated but operationally linked. The retained Figure 2a and "
        "Figure 2b should remain in their original positions and continue to illustrate "
        "the high-level system view and the interaction between perception, digital-twin, "
        "and decision-support layers."
    ),
    (
        "At runtime, the framework uses ROS2-compatible topic interfaces such as "
        "/camera/image_raw, /audio/stream, /robot_pose, /event_log, and /system_health "
        "to structure the data flow. A digital twin mirrors timestamped activity, enables "
        "replay and audit, and now supports predictive validation. On top of that, the "
        "benchmarking layer evaluates deep, classical, hybrid, domain-adapted, and "
        "privacy-preserving emotion-recognition configurations under controlled "
        "experimental conditions."
    ),
    (
        "The mathematical formulation in the next section makes this architecture "
        "explicit as a five-zone information pipeline in which the output of each zone "
        "becomes the input to the next."
    ),
]

MATHEMATICAL_FORMULATION = [
    "## 3. Mathematical Formulation Aligned with the PAEMDT Architecture",
    (
        "This section formalizes the information flow of PAEMDT according to the five "
        "architecture zones. The formulation is intentionally modular and evidence-aware, "
        "so that repository-implemented, simulation-supported, and future clinically "
        "validated components can be distinguished without collapsing them into a single "
        "undifferentiated claim. The output of each zone is defined explicitly and then "
        "used as the input to the next zone, thereby preserving the closed-loop logic of "
        "the PAEMDT framework."
    ),
    "### 3.1 Zone 1 -- Multimodal Perception and Modality Encoding",
    (
        "At time t, the raw multimodal observation is defined as:"
    ),
    "(1) X_t = {x_t^v, x_t^a, x_t^p, x_t^m, x_t^l}",
    (
        "where x_t^v denotes the visual observation, x_t^a denotes the acoustic "
        "observation, x_t^p denotes the physiological observation, x_t^m denotes the "
        "motion or activity observation, and x_t^l denotes the language or contextual "
        "observation."
    ),
    (
        "Each modality is encoded by a modality-specific feature extractor:"
    ),
    "(2) y_t^j = phi_j(x_t^j),   j in {v, a, p, m, l}",
    (
        "The set of modality embeddings is then written as:"
    ),
    "(3) Y_t = {y_t^v, y_t^a, y_t^p, y_t^m, y_t^l}",
    (
        "The output of Zone 1 is Y_t, the set of modality-specific embeddings. This "
        "output becomes the input to Zone 2. In the current repository-validated "
        "baseline, the visual stream is implemented using DeepFace-based facial emotion "
        "analysis, and the acoustic stream is implemented using MFCC-based speech emotion "
        "recognition with an RBF-SVM classifier. Physiological, motion, and language "
        "streams are treated as simulation-supported or planned modules depending on the "
        "current evidence level [10], [15], [18], [26], [27]."
    ),
    "### 3.2 Zone 2 -- Missing-Modality Masking and Cross-Modal Fusion",
    (
        "To model missing or unreliable sensing channels, a modality-availability mask is "
        "defined as:"
    ),
    "(4) M_t = {m_t^v, m_t^a, m_t^p, m_t^m, m_t^l},   m_t^j in {0, 1}",
    (
        "where m_t^j = 1 indicates that modality j is available and m_t^j = 0 indicates "
        "that the modality is missing, corrupted, or judged unreliable."
    ),
    (
        "Fusion over the modality embeddings is defined by a generic operator F(.):"
    ),
    "(5) F_t = F(Y_t, M_t)",
    (
        "In the current repository-aligned implementation, the baseline fusion is "
        "rule-based face-speech fusion. In the extended PAEMDT architecture, F(.) can be "
        "instantiated as an attention-inspired or cross-attention multimodal fusion "
        "operator. A generic attention-style formulation is:"
    ),
    "(6) Q_t = W_Q Y_t,   K_t = W_K Y_t,   V_t = W_V Y_t",
    "(7) F_t = softmax((Q_t K_t^T) / sqrt(d_k)) V_t",
    (
        "The missing-modality mask is then applied to the fused representation:"
    ),
    "(8) F_t^* = F_t odot M_t",
    (
        "The output of Zone 2 is F_t^*, the masked fused multimodal representation. This "
        "output becomes the input to Zone 3 and also supports the reasoning layer in "
        "Zone 4. The repository currently supports a baseline fusion implementation and "
        "benchmark-level attention-inspired fusion logic; full end-to-end cross-attention "
        "learning is treated as an extendable architecture component rather than an "
        "already deployed subsystem [10], [15], [18], [22], [23]."
    ),
    "### 3.3 Zone 3 -- Digital-Twin State Update and Synchronization",
    (
        "The digital-twin state is defined as:"
    ),
    "(9) S_t^DT = {s_t^pat, s_t^rob, s_t^env, s_t^int}",
    (
        "where s_t^pat denotes the patient state, s_t^rob denotes the robot state, "
        "s_t^env denotes the environmental context, and s_t^int denotes the interaction "
        "history."
    ),
    (
        "The twin update is written as:"
    ),
    "(10) S_t^DT = U(S_{t-1}^DT, F_t^*, u_{t-1})",
    (
        "where U(.) is the digital-twin update function and u_{t-1} is the previous "
        "caregiving action fed back from Zone 5 at the previous time step. To measure "
        "synchronization freshness, the digital-twin synchronization error is defined as:"
    ),
    "(11) epsilon_t^DT = | t_now - max(t_t^v, t_t^a, t_t^p, t_t^m, t_t^l) |",
    (
        "This expression measures how far the digital twin lags behind the newest "
        "available multimodal evidence. A large synchronization error indicates stale or "
        "asynchronous evidence and should influence downstream HITL routing. The output "
        "of Zone 3 is therefore the updated digital-twin state S_t^DT together with the "
        "synchronization error epsilon_t^DT. These outputs become inputs to Zone 4 for "
        "explainable reasoning, risk assessment, and safety routing [12], [19], [24], "
        "[25]. In the current validation, the synchronization term is also empirically "
        "supported by a measured mean latency of 124.0 +/- 67.0 ms."
    ),
    "### 3.4 Zone 4 -- Explainability, Privacy-Aware Inference, and HITL Reasoning",
    (
        "PAEMDT represents structured domain knowledge through a care-relevant knowledge "
        "graph:"
    ),
    "(12) G = (V, E)",
    (
        "where V is the set of care-relevant concepts and E is the set of semantic "
        "relations among them."
    ),
    (
        "Knowledge-grounded evidence retrieval is defined as:"
    ),
    "(13) E_t = R(G, F_t^*, S_t^DT)",
    (
        "where E_t is the set of retrieved evidence nodes and relations associated with "
        "the current fused observation and the current digital-twin state."
    ),
    (
        "Explanation generation is defined as:"
    ),
    "(14) e_t = G_exp(F_t^*, S_t^DT, E_t)",
    (
        "To support downstream reasoning, the fused representation, twin state, and "
        "retrieved evidence are combined into a single reasoning vector:"
    ),
    "(15) z_t = [F_t^*, S_t^DT, E_t]",
    (
        "The health-risk score is defined as:"
    ),
    "(16) r_t = sigma(w_r^T z_t + b_r)",
    (
        "and the anomaly score is defined as:"
    ),
    "(17) a_t = D(z_t, z_ref)",
    (
        "where D(.) is a deviation measure with respect to a reference baseline state "
        "z_ref."
    ),
    (
        "Privacy-aware transformation over the raw evidence is represented as:"
    ),
    "(18) X_t_tilde = Pi(X_t; lambda)",
    (
        "where Pi(.) is the privacy-control function and lambda is the selected privacy "
        "configuration. In the current repository-aligned interpretation, Pi(.) "
        "represents architecture-level privacy control and privacy-utility analysis, not "
        "a formal differential-privacy guarantee by itself. Formal privacy guarantees are "
        "reported only when epsilon accounting is explicitly implemented in the training "
        "pipeline."
    ),
    (
        "The HITL routing tier is then defined as:"
    ),
    "(19) T_t = { urgent escalation, if r_t > tau_r^high OR a_t > tau_a^high; caregiver review, if r_t > tau_r^mid OR a_t > tau_a^mid; autonomous response, otherwise }",
    (
        "The output of Zone 4 is the tuple {e_t, r_t, a_t, X_t_tilde, T_t}. These outputs "
        "become inputs to Zone 5 for tiered caregiving action selection. This zone "
        "therefore couples explainability, risk scoring, anomaly assessment, privacy "
        "filtering, and HITL escalation within one reasoning layer [11], [13], [14], "
        "[16], [17], [20], [21], [22], [23], [28], [29]."
    ),
    "### 3.5 Zone 5 -- Tiered Caregiving Action and Deployment-Oriented Model Selection",
    (
        "The caregiving action policy is defined as:"
    ),
    "(20) u_t = pi(F_t^*, S_t^DT, e_t, r_t, a_t, X_t_tilde, T_t)",
    (
        "where u_t belongs to the set {autonomous response, caregiver review, urgent "
        "escalation}. Zone 5 closes the PAEMDT loop. The selected action u_t is logged "
        "and becomes part of the next digital-twin update through u_{t-1} in Zone 3 at "
        "the following time step."
    ),
    (
        "Deployment-oriented model selection is represented by:"
    ),
    "(21) Score_k = alpha F1_k + beta R_k + gamma C_k - delta L_k - eta P_k",
    (
        "where F1_k is the macro-F1 score of model k, R_k is external-domain robustness, "
        "C_k is calibration quality, L_k is latency, and P_k is privacy cost. The "
        "numerical values of these weights are specified only in the case-study section "
        "because they are experiment-specific and should not be treated as universal "
        "model parameters."
    ),
    (
        "External-domain robustness is defined as:"
    ),
    "(22) R_k = Acc_k^ext / Acc_k^val",
    (
        "and calibration error is defined as:"
    ),
    "(23) ECE = sum_{b=1}^{B} (|B_b| / N) | acc(B_b) - conf(B_b) |",
    (
        "where B_b is the set of samples in calibration bin b, N is the total number of "
        "samples, acc(B_b) is the empirical accuracy in bin b, and conf(B_b) is the mean "
        "confidence in bin b. The model-selection score prevents the system from "
        "selecting a model based only on internal benchmark accuracy. This is important "
        "because a model may have high validation accuracy but weak external robustness, "
        "poor calibration, excessive latency, or unfavorable privacy characteristics "
        "[22], [23], [25], [26], [27], [28]."
    ),
    (
        "The formulation above defines the architecture-level information flow of "
        "PAEMDT. The following case study evaluates which parts of this pipeline are "
        "repository-implemented, simulation-supported, or planned for future clinical "
        "validation."
    ),
]

TABLE_SPECS = {
    "4.1": [
        (
            "Table 4. Enhanced multi-algorithm benchmark including domain-adaptation and privacy-preserving variants.",
            PROJECT_ROOT / "experiments" / "results" / "paper_tables" / "table4_multi_algorithm_benchmark.csv",
        ),
        (
            "Table 4b. Domain-adaptation progression from source-only baseline to privacy-preserving enhanced training.",
            PROJECT_ROOT / "experiments" / "results" / "paper_tables" / "table_domain_adaptation_results.csv",
        ),
    ],
    "4.6": [
        (
            "Table 5. Component-wise ablation analysis of the PAEMDT framework.",
            PROJECT_ROOT / "experiments" / "results" / "paper_tables" / "table5_ablation.csv",
        ),
    ],
}


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    paragraph.paragraph_format.space_after = Pt(0)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(11)

    for name, size in [("Title", 16), ("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 11)]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True


def add_title_block(doc: Document) -> None:
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(TITLE)

    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors.add_run(AUTHORS).bold = True

    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def add_paragraphs(doc: Document, paragraphs: list[str]) -> None:
    for text in paragraphs:
        paragraph = doc.add_paragraph(text)
        paragraph.paragraph_format.space_after = Pt(6)


def add_structured_blocks(doc: Document, blocks: list[str]) -> None:
    for item in blocks:
        if item.startswith("## "):
            add_heading(doc, item[3:].strip(), level=1)
        elif item.startswith("### "):
            add_heading(doc, item[4:].strip(), level=2)
        else:
            paragraph = doc.add_paragraph(item)
            paragraph.paragraph_format.space_after = Pt(6)


def read_companion_lines() -> list[str]:
    lines = COMPANION_MD.read_text(encoding="utf-8").splitlines()
    start = 0
    for i, line in enumerate(lines):
        if line.strip() == "## 4. Case Study":
            start = i
            break
    return lines[start:]


def parse_blocks(lines: list[str]) -> list[tuple[str, str]]:
    blocks: list[tuple[str, str]] = []
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        if paragraph_buffer:
            text = " ".join(part.strip() for part in paragraph_buffer if part.strip())
            if text:
                blocks.append(("paragraph", text))
            paragraph_buffer.clear()

    image_pattern = re.compile(r"!\[(?P<alt>.*?)\]\((?P<path>.*?)\)")

    for raw_line in lines:
        line = raw_line.rstrip()
        if not line.strip():
            flush_paragraph()
            continue
        if line.startswith("#"):
            flush_paragraph()
            level = len(line) - len(line.lstrip("#"))
            text = line[level:].strip()
            blocks.append((f"heading{level}", text))
            continue
        match = image_pattern.match(line.strip())
        if match:
            flush_paragraph()
            blocks.append(("image", match.group("path")))
            continue
        if line.startswith("**Figure "):
            flush_paragraph()
            caption_text = re.sub(r"\*\*(Figure\s+\d+\.?)\*\*", r"\1", line).strip()
            blocks.append(("caption", caption_text))
            continue
        paragraph_buffer.append(line)

    flush_paragraph()
    return blocks


def add_csv_table(doc: Document, caption: str, csv_path: Path) -> None:
    if not csv_path.exists():
        return

    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = caption_p.add_run(caption)
    cap_run.bold = True

    with csv_path.open("r", encoding="utf-8-sig", newline="") as handle:
        reader = list(csv.reader(handle))
    if not reader:
        return

    headers = reader[0]
    rows = reader[1:]

    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True

    for col, header in enumerate(headers):
        set_cell_text(table.cell(0, col), header, bold=True)

    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            set_cell_text(table.cell(row_idx, col_idx), value)

    doc.add_paragraph()


def add_image(doc: Document, image_rel_path: str) -> None:
    image_path = (COMPANION_MD.parent / image_rel_path).resolve()
    if not image_path.exists():
        return
    doc.add_picture(str(image_path), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def write_complete_markdown(lines_4_6: list[str]) -> None:
    output_lines: list[str] = [
        f"# {TITLE}",
        "",
        f"**Author:** {AUTHORS}",
        "",
        "## Abstract",
        "",
        ABSTRACT,
        "",
        "## 1. Introduction",
        "",
    ]
    output_lines.extend([para for paragraph in INTRODUCTION for para in [paragraph, ""]])
    output_lines.append("## 2. Framework Architecture")
    output_lines.append("")
    output_lines.extend([para for paragraph in FRAMEWORK_ARCHITECTURE for para in [paragraph, ""]])
    for item in MATHEMATICAL_FORMULATION:
        output_lines.append(item)
        output_lines.append("")
    output_lines.extend(lines_4_6)
    OUTPUT_MD.write_text("\n".join(output_lines).strip() + "\n", encoding="utf-8")


def build_docx() -> None:
    lines_4_6 = read_companion_lines()
    write_complete_markdown(lines_4_6)

    doc = Document()
    style_document(doc)
    add_title_block(doc)

    add_heading(doc, "Abstract", level=1)
    add_paragraphs(doc, [ABSTRACT])

    add_heading(doc, "1. Introduction", level=1)
    add_paragraphs(doc, INTRODUCTION)

    add_heading(doc, "2. Framework Architecture", level=1)
    add_paragraphs(doc, FRAMEWORK_ARCHITECTURE)

    add_structured_blocks(doc, MATHEMATICAL_FORMULATION)

    blocks = parse_blocks(lines_4_6)
    current_subsection: str | None = None
    inserted_tables: set[str] = set()

    for block_type, payload in blocks:
        if block_type.startswith("heading"):
            heading_level = int(block_type.replace("heading", ""))
            if current_subsection and current_subsection in TABLE_SPECS and current_subsection not in inserted_tables:
                for caption, csv_path in TABLE_SPECS[current_subsection]:
                    add_csv_table(doc, caption, csv_path)
                inserted_tables.add(current_subsection)

            if heading_level == 2:
                add_heading(doc, payload, level=1)
            elif heading_level == 3:
                add_heading(doc, payload, level=2)
                subsection_match = re.match(r"^(\d+\.\d+)", payload)
                current_subsection = subsection_match.group(1) if subsection_match else None
            else:
                add_heading(doc, payload, level=min(heading_level, 3))
        elif block_type == "paragraph":
            paragraph = doc.add_paragraph(payload)
            paragraph.paragraph_format.space_after = Pt(6)
        elif block_type == "image":
            add_image(doc, payload)
        elif block_type == "caption":
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = caption.add_run(payload)
            run.italic = True
            caption.paragraph_format.space_after = Pt(8)

    if current_subsection and current_subsection in TABLE_SPECS and current_subsection not in inserted_tables:
        for caption, csv_path in TABLE_SPECS[current_subsection]:
            add_csv_table(doc, caption, csv_path)

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)


if __name__ == "__main__":
    build_docx()
    print(f"Wrote {OUTPUT_DOCX}")
    print(f"Wrote {OUTPUT_MD}")
