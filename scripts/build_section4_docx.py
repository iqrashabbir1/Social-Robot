from __future__ import annotations

import csv
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
SECTION4_MD = PROJECT_ROOT / "docs" / "paper1" / "final_package" / "section4_experimental_evaluation_elsevier.md"
OUTPUT_DOCX = PROJECT_ROOT / "docs" / "paper1" / "final_package" / "section4_case_study_complete_with_figures_repo_aligned.docx"

TABLE_INSERTS = {
    "4.1": [
        (
            "Table 4. Enhanced benchmark comparison including domain-adaptation and privacy-preserving variants.",
            PROJECT_ROOT / "experiments" / "results" / "paper_tables" / "table4_multi_algorithm_benchmark.csv",
        ),
        (
            "Table 4b. Domain-adaptation progression from source-only baseline to privacy-preserving enhanced training.",
            PROJECT_ROOT / "experiments" / "results" / "paper_tables" / "table_domain_adaptation_results.csv",
        ),
    ],
    "4.5": [
        (
            "Table 5. Component-wise ablation analysis of the PAEMDT framework.",
            PROJECT_ROOT / "experiments" / "results" / "paper_tables" / "table5_ablation.csv",
        ),
    ],
    "4.8": [
        (
            "Table 6. Missing-modality robustness and escalation outcomes across degraded sensing conditions.",
            PROJECT_ROOT / "outputs" / "tables" / "paper1_table_missing_modality_robustness.csv",
        ),
    ],
    "4.9": [
        (
            "Table 7. Edge hardware benchmark results across currently tracked deployment platforms.",
            PROJECT_ROOT / "outputs" / "tables" / "paper1_table_edge_benchmark.csv",
        ),
    ],
}


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(11)

    for name, size in [("Title", 16), ("Heading 1", 14), ("Heading 2", 12)]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    p.paragraph_format.space_after = Pt(0)


def add_csv_table(doc: Document, caption: str, csv_path: Path) -> None:
    if not csv_path.exists():
        return
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_p.add_run(caption)
    caption_run.bold = True

    with csv_path.open("r", encoding="utf-8-sig", newline="") as handle:
        reader = list(csv.reader(handle))
    if not reader:
        return

    headers = reader[0]
    rows = reader[1:]
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for idx, header in enumerate(headers):
        set_cell_text(table.cell(0, idx), header, bold=True)
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            set_cell_text(table.cell(row_idx, col_idx), value)
    doc.add_paragraph()


def parse_markdown(path: Path) -> list[tuple[str, str]]:
    lines = path.read_text(encoding="utf-8").splitlines()
    blocks: list[tuple[str, str]] = []
    paragraph_buffer: list[str] = []
    image_pattern = re.compile(r"!\[(?P<alt>.*?)\]\((?P<path>.*?)\)")

    def flush() -> None:
        if paragraph_buffer:
            blocks.append(("paragraph", " ".join(part.strip() for part in paragraph_buffer if part.strip())))
            paragraph_buffer.clear()

    for raw_line in lines:
        line = raw_line.rstrip()
        if not line.strip():
            flush()
            continue
        if line.startswith("# "):
            flush()
            blocks.append(("heading1", line[2:].strip()))
            continue
        if line.startswith("## "):
            flush()
            blocks.append(("heading2", line[3:].strip()))
            continue
        match = image_pattern.match(line.strip())
        if match:
            flush()
            blocks.append(("image", match.group("path")))
            continue
        paragraph_buffer.append(line)
    flush()
    return blocks


def build() -> None:
    doc = Document()
    style_document(doc)

    blocks = parse_markdown(SECTION4_MD)
    current_subsection: str | None = None
    inserted: set[str] = set()

    for block_type, payload in blocks:
        if block_type == "heading1":
            doc.add_paragraph(payload, style="Title")
        elif block_type == "heading2":
            subsection_match = re.match(r"^(\d+\.\d+)", payload)
            if current_subsection and current_subsection in TABLE_INSERTS and current_subsection not in inserted:
                for caption, csv_path in TABLE_INSERTS[current_subsection]:
                    add_csv_table(doc, caption, csv_path)
                inserted.add(current_subsection)
            doc.add_heading(payload, level=1)
            current_subsection = subsection_match.group(1) if subsection_match else None
        elif block_type == "image":
            image_path = (SECTION4_MD.parent / payload).resolve()
            if image_path.exists():
                doc.add_picture(str(image_path), width=Inches(6.2))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p = doc.add_paragraph(payload)
            p.paragraph_format.space_after = Pt(6)

    if current_subsection and current_subsection in TABLE_INSERTS and current_subsection not in inserted:
        for caption, csv_path in TABLE_INSERTS[current_subsection]:
            add_csv_table(doc, caption, csv_path)

    doc.save(OUTPUT_DOCX)
    print(f"Wrote {OUTPUT_DOCX}")


if __name__ == "__main__":
    build()
